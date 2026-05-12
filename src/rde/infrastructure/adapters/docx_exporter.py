"""DocxExporter — Adapter implementing DocumentExporterPort.

Exports EDAReport to Word (.docx) and PDF formats.
Follows medpaper's template-based pattern:
  EDAReport → structured document with embedded figures & tables.

Uses python-docx for .docx generation.
PDF requires optional xhtml2pdf; falls back to HTML if unavailable.
"""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Any

from rde.domain.models.report import EDAReport, ReportSection
from rde.domain.ports import DocumentExporterPort

logger = logging.getLogger(__name__)


class DocxExporter(DocumentExporterPort):
    """Export EDAReport to .docx and .pdf formats."""

    # ── Style constants ──────────────────────────────────────────────
    FONT_NAME = "Times New Roman"
    FONT_SIZE_BODY = 11
    FONT_SIZE_TITLE = 16
    FIGURE_WIDTH_INCHES = 5.5
    TABLE_STYLE = "Table Grid"

    def export_docx(
        self,
        report: Any,
        output_path: Path,
        *,
        figures_dir: Path | None = None,
    ) -> Path:
        """Export EDAReport as Word .docx with embedded figures and tables."""
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        rpt: EDAReport = report
        doc = Document()

        # ── Document-level font defaults ─────────────────────────────
        style = doc.styles["Normal"]
        font = style.font
        font.name = self.FONT_NAME
        font.size = Pt(self.FONT_SIZE_BODY)

        # ── Title page ───────────────────────────────────────────────
        title_para = doc.add_heading(rpt.title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        formal_profile = rpt.metadata.get("_export_profile") == "formal_research_report"
        meta_lines = (
            []
            if formal_profile
            else [
                f"Dataset: {rpt.dataset_id}",
                f"Project: {rpt.project_id}",
                f"Generated: {rpt.created_at.strftime('%Y-%m-%d %H:%M')}",
            ]
        )
        for line in meta_lines:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

        # ── Sections ─────────────────────────────────────────────────
        for section in rpt.sections:
            self._render_section(doc, section, figures_dir)

        # ── Metadata footer ──────────────────────────────────────────
        public_metadata = {k: v for k, v in rpt.metadata.items() if not str(k).startswith("_")}
        if public_metadata:
            doc.add_heading("Metadata", level=2)
            for k, v in public_metadata.items():
                doc.add_paragraph(f"{k}: {v}", style="List Bullet")

        # ── Save ─────────────────────────────────────────────────────
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        return output_path

    def export_pdf(
        self,
        report: Any,
        output_path: Path,
        *,
        figures_dir: Path | None = None,
    ) -> Path:
        """Export EDAReport as PDF via HTML rendering.

        Uses xhtml2pdf (pure Python, no native deps).
        Falls back to HTML file if xhtml2pdf unavailable.
        """
        rpt: EDAReport = report
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        # Build HTML with embedded figures (base64)
        html_content = self._build_pdf_html(rpt, figures_dir)

        try:
            from xhtml2pdf import pisa  # type: ignore[import-untyped]

            with open(str(output_path), "wb") as f:
                status = pisa.CreatePDF(html_content, dest=f, encoding="utf-8")
            if status.err:
                raise ValueError(f"xhtml2pdf 轉換失敗，錯誤數: {status.err}")
            return output_path
        except ImportError:
            # Fallback: save HTML so user can open in browser → Print to PDF
            html_path = output_path.with_suffix(".html")
            html_path.write_text(html_content, encoding="utf-8")
            logger.warning(
                "xhtml2pdf unavailable; exported HTML fallback instead of PDF: %s",
                html_path,
            )
            return html_path

    # ── Private helpers ──────────────────────────────────────────────

    def _render_section(
        self,
        doc: Any,
        section: ReportSection,
        figures_dir: Path | None,
    ) -> None:
        """Render a single ReportSection into the Word document."""
        doc.add_heading(section.title, level=2)

        # ── Content (parse markdown-like text) ───────────────────────
        self._add_markdown_content(doc, section.content, figures_dir)

        # ── Tables ───────────────────────────────────────────────────
        for tbl_data in section.tables:
            self._add_table(doc, tbl_data)

        # ── Figures ──────────────────────────────────────────────────
        for fig_path_str in section.figures:
            fig_path = Path(fig_path_str)
            if not fig_path.is_absolute() and figures_dir:
                fig_path = figures_dir / fig_path.name
            if fig_path.exists():
                self._add_picture_with_caption(
                    doc,
                    fig_path,
                    f"Figure: {fig_path.stem}",
                    width_inches=self.FIGURE_WIDTH_INCHES,
                )

    def _add_markdown_content(
        self, doc: Any, content: str, figures_dir: Path | None = None
    ) -> None:
        """Parse markdown-like content and add to document.

        Handles: paragraphs, bold (**text**), bullet lists (- item),
        markdown tables (| col | col |), subheadings (# ## ### text),
        blockquotes (> text), horizontal rules (---), figure references,
        fenced code blocks (```).
        """
        from docx.shared import Pt, RGBColor

        lines = content.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # Empty line — skip
            if not stripped:
                i += 1
                continue

            # Fenced code block (```)
            if stripped.startswith("```"):
                i += 1
                code_lines: list[str] = []
                while i < len(lines):
                    if lines[i].strip().startswith("```"):
                        i += 1
                        break
                    code_lines.append(lines[i].rstrip())
                    i += 1
                if code_lines:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Pt(18)
                    run = p.add_run("\n".join(code_lines))
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                continue

            # Headings (# ## ### ####)
            if stripped.startswith("#### "):
                doc.add_heading(stripped[5:].strip(), level=4)
                i += 1
                continue
            if stripped.startswith("### "):
                doc.add_heading(stripped[4:].strip(), level=3)
                i += 1
                continue
            if stripped.startswith("## "):
                doc.add_heading(stripped[3:].strip(), level=2)
                i += 1
                continue
            if stripped.startswith("# "):
                doc.add_heading(stripped[2:].strip(), level=1)
                i += 1
                continue

            # Horizontal rule
            if re.match(r"^-{3,}$", stripped) or re.match(r"^\*{3,}$", stripped):
                doc.add_paragraph("─" * 50)
                i += 1
                continue

            # Markdown image reference — check BEFORE blockquote.
            line_for_fig = stripped
            if line_for_fig.startswith("> "):
                line_for_fig = line_for_fig[2:].strip()
            image_match = self._match_markdown_image(line_for_fig)
            if image_match and figures_dir:
                alt_text, target = image_match
                fig_path = self._resolve_figure_path(target, figures_dir)
                if fig_path and fig_path.exists():
                    self._add_picture_with_caption(
                        doc,
                        fig_path,
                        alt_text or f"Figure: {fig_path.stem}",
                        width_inches=self.FIGURE_WIDTH_INCHES,
                    )
                else:
                    p = doc.add_paragraph()
                    self._add_formatted_run(p, line_for_fig)
                i += 1
                continue

            # Legacy figure reference — check BEFORE blockquote.
            line_for_fig_clean = self._strip_md_formatting(line_for_fig)
            fig_match = re.search(
                r"\[.*?(?:Figure|Fig|圖)\s*\d*[:\s]+([^\]]+\.png)\]",
                line_for_fig_clean,
            )
            if fig_match and figures_dir:
                fig_name = fig_match.group(1).strip()
                fig_path = self._resolve_figure_path(fig_name, figures_dir)
                if fig_path and fig_path.exists():
                    self._add_picture_with_caption(
                        doc,
                        fig_path,
                        f"Figure: {fig_path.stem}",
                        width_inches=self.FIGURE_WIDTH_INCHES,
                    )
                else:
                    p = doc.add_paragraph()
                    self._add_formatted_run(p, line_for_fig)
                i += 1
                continue

            # Blockquote (> text) — preserve bold/italic formatting
            if stripped.startswith("> "):
                quote_text = stripped[2:].strip()
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(36)
                self._add_formatted_run(p, quote_text)
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                i += 1
                continue

            # Bullet list
            if stripped.startswith("- ") or stripped.startswith("* "):
                doc.add_paragraph(
                    self._strip_md_formatting(stripped[2:].strip()),
                    style="List Bullet",
                )
                i += 1
                continue

            # Numbered list
            if re.match(r"^\d+\.\s", stripped):
                text = re.sub(r"^\d+\.\s", "", stripped).strip()
                doc.add_paragraph(
                    self._strip_md_formatting(text),
                    style="List Number",
                )
                i += 1
                continue

            # Markdown table
            if "|" in stripped and stripped.startswith("|"):
                table_lines = []
                while i < len(lines) and "|" in lines[i] and lines[i].strip().startswith("|"):
                    table_lines.append(lines[i])
                    i += 1
                self._add_md_table(doc, table_lines)
                continue

            # Regular paragraph
            if stripped:
                p = doc.add_paragraph()
                self._add_formatted_run(p, stripped)

            i += 1

    def _add_formatted_run(self, paragraph: Any, text: str) -> None:
        """Add a run with basic markdown formatting (bold, italic)."""
        # Split by **bold** and *italic* markers
        parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*"):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            else:
                paragraph.add_run(part)

    @staticmethod
    def _strip_md_formatting(text: str) -> str:
        """Remove markdown bold/italic markers for plain text contexts."""
        text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
        text = re.sub(r"\*([^*]+)\*", r"\1", text)
        return text

    def _add_table(self, doc: Any, tbl_data: dict[str, Any]) -> None:
        """Add a structured dict table to the Word document."""
        if not tbl_data:
            return
        if isinstance(tbl_data.get("headers"), list) and isinstance(tbl_data.get("rows"), list):
            headers = [str(value) for value in tbl_data.get("headers", [])]
            rows_data = [list(row) for row in tbl_data.get("rows", [])]
            if not headers:
                return

            table = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
            table.style = self.TABLE_STYLE

            for j, header in enumerate(headers):
                table.rows[0].cells[j].text = header

            for i, row in enumerate(rows_data):
                padded = [str(value) for value in row] + [""] * (len(headers) - len(row))
                for j, value in enumerate(padded[: len(headers)]):
                    table.rows[i + 1].cells[j].text = value

            doc.add_paragraph()
            return

        headers = list(tbl_data.keys())
        first_vals = tbl_data[headers[0]]
        if isinstance(first_vals, dict):
            row_keys = list(first_vals.keys())
        else:
            doc.add_paragraph(str(tbl_data))
            return

        table = doc.add_table(rows=1 + len(row_keys), cols=len(headers))
        table.style = self.TABLE_STYLE

        # Header row
        for j, h in enumerate(headers):
            table.rows[0].cells[j].text = str(h)

        # Data rows
        for i, rk in enumerate(row_keys):
            for j, h in enumerate(headers):
                val = tbl_data[h].get(rk, "") if isinstance(tbl_data[h], dict) else tbl_data[h]
                table.rows[i + 1].cells[j].text = str(val)

        doc.add_paragraph()  # spacing after table

    def _add_md_table(self, doc: Any, table_lines: list[str]) -> None:
        """Parse markdown table lines (| col | col |) into a Word table."""
        rows_data: list[list[str]] = []
        for line in table_lines:
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            # Skip separator row (--- | ---)
            if all(re.match(r"^[-:]+$", c) for c in cells):
                continue
            rows_data.append(cells)

        if not rows_data:
            return

        n_cols = max(len(r) for r in rows_data)
        table = doc.add_table(rows=len(rows_data), cols=n_cols)
        table.style = self.TABLE_STYLE

        for i, row in enumerate(rows_data):
            for j, cell_text in enumerate(row):
                if j < n_cols:
                    table.rows[i].cells[j].text = self._strip_md_formatting(cell_text)

        doc.add_paragraph()  # spacing

    def _add_picture_with_caption(
        self,
        doc: Any,
        fig_path: Path,
        caption: str,
        *,
        width_inches: float,
    ) -> None:
        """Insert a Word-safe figure with a centered caption."""
        from docx.shared import Inches, Pt, RGBColor

        p = doc.add_paragraph()
        p.paragraph_format.alignment = 1  # CENTER
        run = p.add_run()
        image_source = self._docx_ready_image_source(fig_path)
        if hasattr(image_source, "seek"):
            image_source.seek(0)
        run.add_picture(image_source, width=Inches(width_inches))

        cap = doc.add_paragraph(caption)
        cap_fmt = cap.paragraph_format
        cap_fmt.alignment = 1  # CENTER
        for run in cap.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    def _docx_ready_image_source(self, fig_path: Path) -> Any:
        """Return an RGB PNG stream so Word/PDF renderers avoid alpha issues."""
        import io

        try:
            from PIL import Image
        except ImportError:
            return str(fig_path)

        try:
            with Image.open(fig_path) as image:
                rgba = image.convert("RGBA")
                background = Image.new("RGB", rgba.size, "white")
                background.paste(rgba, mask=rgba.getchannel("A"))
                buffer = io.BytesIO()
                background.save(buffer, format="PNG", optimize=True)
                buffer.seek(0)
                return buffer
        except Exception:
            logger.exception("Failed to normalize image for DOCX export: %s", fig_path)
            return str(fig_path)

    def _build_pdf_html(
        self,
        report: EDAReport,
        figures_dir: Path | None = None,
    ) -> str:
        """Build a self-contained HTML for PDF rendering.

        Embeds figures as base64 data URIs so the HTML/PDF is portable.
        """
        lines = [
            "<!DOCTYPE html>",
            '<html lang="zh-TW">',
            "<head>",
            f"<title>{report.title}</title>",
            '<meta charset="utf-8">',
            "<style>",
            "  @page { size: A4; margin: 2cm; }",
            "  body { font-family: 'Times New Roman', serif; font-size: 11pt;",
            "         max-width: 700px; margin: auto; line-height: 1.6; }",
            "  h1 { text-align: center; font-size: 18pt; margin-bottom: 0.5em; }",
            "  h2 { font-size: 14pt; border-bottom: 1px solid #ccc;",
            "       padding-bottom: 4px; margin-top: 1.5em; }",
            "  h3 { font-size: 12pt; }",
            "  table { border-collapse: collapse; width: 100%; margin: 1em 0;",
            "          font-size: 10pt; }",
            "  th, td { border: 1px solid #999; padding: 6px 8px; text-align: left; }",
            "  th { background: #f0f0f0; font-weight: bold; }",
            "  img { max-width: 100%; display: block; margin: 1em auto; }",
            "  .caption { text-align: center; font-size: 10pt; color: #555;",
            "             margin-top: -0.5em; margin-bottom: 1em; }",
            "  .metadata { font-size: 10pt; color: #666; text-align: center; }",
            "  hr { margin: 2em 0; border: none; border-top: 1px solid #ccc; }",
            "</style>",
            "</head>",
            "<body>",
            f"<h1>{report.title}</h1>",
            '<div class="metadata">',
            ""
            if report.metadata.get("_export_profile") == "formal_research_report"
            else (
                f"<p>Dataset: {report.dataset_id} | "
                f"Project: {report.project_id} | "
                f"Generated: {report.created_at.strftime('%Y-%m-%d %H:%M')}</p>"
            ),
            "</div>",
            "<hr>",
        ]

        for section in report.sections:
            lines.append(f"<h2>{section.title}</h2>")

            # Content — convert basic markdown to HTML
            for para in section.content.split("\n"):
                para = para.strip()
                if not para:
                    continue
                image_match = self._match_markdown_image(para)
                if image_match and figures_dir:
                    alt_text, target = image_match
                    fig_path = self._resolve_figure_path(target, figures_dir)
                    if fig_path and fig_path.exists():
                        lines.append(self._figure_to_html(fig_path, alt_text))
                    else:
                        lines.append(f"<p>{self._md_to_html_inline(para)}</p>")
                    continue
                if para.startswith("### "):
                    lines.append(f"<h3>{para[4:]}</h3>")
                elif para.startswith("- ") or para.startswith("* "):
                    lines.append(f"<li>{self._md_to_html_inline(para[2:])}</li>")
                elif para.startswith("|") and "|" in para:
                    # Let markdown table lines through — collected below
                    lines.append(f"<p>{para}</p>")
                else:
                    lines.append(f"<p>{self._md_to_html_inline(para)}</p>")

            # Tables
            if section.tables:
                for tbl in section.tables:
                    lines.append(self._dict_to_html_table(tbl))

            # Figures — embed as base64
            for fig_path_str in section.figures:
                fig_path = Path(fig_path_str)
                if not fig_path.is_absolute() and figures_dir:
                    fig_path = figures_dir / fig_path.name
                if fig_path.exists():
                    lines.append(self._figure_to_html(fig_path, f"Figure: {fig_path.stem}"))

        public_metadata = {k: v for k, v in report.metadata.items() if not str(k).startswith("_")}
        if public_metadata:
            lines.append("<h2>Metadata</h2>")
            lines.append("<ul>")
            for k, v in public_metadata.items():
                lines.append(f"<li><strong>{k}:</strong> {v}</li>")
            lines.append("</ul>")

        lines.extend(["</body>", "</html>"])
        return "\n".join(lines)

    @staticmethod
    def _md_to_html_inline(text: str) -> str:
        """Convert inline markdown (bold, italic) to HTML tags."""
        text = re.sub(r"\*\*([^*]+)\*\*", r"<strong>\1</strong>", text)
        text = re.sub(r"\*([^*]+)\*", r"<em>\1</em>", text)
        return text

    @staticmethod
    def _match_markdown_image(text: str) -> tuple[str, str] | None:
        match = re.match(r"^!\[([^\]]*)\]\(([^)\s]+)(?:\s+\"[^\"]*\")?\)\s*$", text.strip())
        if not match:
            return None
        return match.group(1).strip(), match.group(2).strip("<>")

    @staticmethod
    def _resolve_figure_path(target: str, figures_dir: Path | None) -> Path | None:
        if not target:
            return None
        candidate = Path(target)
        if candidate.is_absolute():
            return candidate
        if figures_dir is None:
            return candidate
        direct = figures_dir / candidate
        if direct.exists():
            return direct
        return figures_dir / candidate.name

    @staticmethod
    def _figure_to_html(fig_path: Path, caption: str) -> str:
        import base64

        data = base64.b64encode(fig_path.read_bytes()).decode()
        suffix = fig_path.suffix.lstrip(".").lower()
        mime = f"image/{suffix}" if suffix != "jpg" else "image/jpeg"
        safe_caption = DocxExporter._md_to_html_inline(caption)
        return f'<img src="data:{mime};base64,{data}">\n<p class="caption">{safe_caption}</p>'

    @staticmethod
    def _dict_to_html_table(tbl: dict[str, Any]) -> str:
        """Convert a dict-of-dicts table to HTML table string."""
        if not tbl:
            return ""
        if isinstance(tbl.get("headers"), list) and isinstance(tbl.get("rows"), list):
            headers = [str(value) for value in tbl.get("headers", [])]
            rows = [list(row) for row in tbl.get("rows", [])]
            if not headers:
                return ""

            lines = ["<table>", "<tr>"]
            for header in headers:
                lines.append(f"<th>{header}</th>")
            lines.append("</tr>")
            for row in rows:
                padded = [str(value) for value in row] + [""] * (len(headers) - len(row))
                lines.append("<tr>")
                for value in padded[: len(headers)]:
                    lines.append(f"<td>{value}</td>")
                lines.append("</tr>")
            lines.append("</table>")
            return "\n".join(lines)

        headers = list(tbl.keys())
        first_vals = tbl[headers[0]]
        if isinstance(first_vals, dict):
            row_keys = list(first_vals.keys())
        else:
            return f"<pre>{tbl}</pre>"

        lines = ["<table>", "<tr>"]
        for h in headers:
            lines.append(f"<th>{h}</th>")
        lines.append("</tr>")
        for rk in row_keys:
            lines.append("<tr>")
            for h in headers:
                val = tbl[h].get(rk, "") if isinstance(tbl[h], dict) else tbl[h]
                lines.append(f"<td>{val}</td>")
            lines.append("</tr>")
        lines.append("</table>")
        return "\n".join(lines)
